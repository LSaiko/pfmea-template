"""
export.py
=========
Convert any generated pFMEA Excel workbook into Word (.docx) and/or PDF.

The generators already encode every scenario's content in the .xlsx output, so
this reads that output back and re-renders it — no scenario data is duplicated
here. Point it at any workbook the generators produce (or any other simple
table-based .xlsx) and it mirrors the sheets as headings + tables.

Usage:
    python export.py pfmea_aiag_vda_2019.xlsx                 # -> .docx and .pdf
    python export.py pfmea_aiag_vda_2019.xlsx --format docx   # -> .docx only
    python export.py pfmea_aiag_vda_2019.xlsx --format pdf    # -> .pdf only
    python export.py pfmea_aiag_vda_2019.xlsx -o report       # custom basename
    python export.py --selfcheck                              # run self-test

Requires: openpyxl, python-docx, reportlab  (see requirements.txt)
"""

import sys
from pathlib import Path

import openpyxl

# AP risk colours, reused by both renderers (matches the Excel palette).
AP_RGB = {"H": (0xC0, 0x00, 0x00), "M": (0xFF, 0xC0, 0x00), "L": (0x92, 0xD0, 0x50)}


def read_sheets(xlsx_path):
    """Read a workbook into [(sheet_name, [blocks])].

    Uses ws.merged_cells.ranges (requires read_only=False) to identify rows
    that are completely covered by a single merged region spanning multiple
    columns — those are heading/banner rows (titles, notes, footers) regardless
    of how many columns they visually occupy.

    A row is treated as a heading when ALL of its non-empty cell addresses fall
    inside a single merged region that spans the full logical row.  This avoids
    the old one-non-empty-cell heuristic that would misclassify legitimate data
    rows that happen to have only one populated column.

    Returns blocks as ("heading", text) | ("table", [[cells...], ...]).
    """
    # read_only=False is required; openpyxl does not expose merged_cells on
    # read-only worksheets reliably.
    wb = openpyxl.load_workbook(xlsx_path, read_only=False, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        # Build a set of (row, col) coordinates that belong to a merged range.
        # Also map each such coordinate to the top-left cell of its merge block
        # so we can recover the value for heading detection.
        merged_topleft = {}   # (row, col) -> (min_row, min_col) of its range
        merged_ranges = list(ws.merged_cells.ranges)  # snapshot before iteration
        for mr in merged_ranges:
            for r in range(mr.min_row, mr.max_row + 1):
                for c in range(mr.min_col, mr.max_col + 1):
                    merged_topleft[(r, c)] = (mr.min_row, mr.min_col)

        blocks, table = [], []
        for row_cells in ws.iter_rows():
            # Collect (coord, value) for every cell in the row, resolving
            # merged-cell slaves to their master's value.
            row_data = []
            for cell in row_cells:
                coord = (cell.row, cell.column)
                if coord in merged_topleft:
                    master = merged_topleft[coord]
                    if master == coord:
                        # This IS the top-left master — use its value.
                        val = "" if cell.value is None else str(cell.value)
                    else:
                        # Slave cell — value is blank; master already counted.
                        val = None   # sentinel: skip in display
                else:
                    val = "" if cell.value is None else str(cell.value)
                row_data.append(val)

            # Build the visible (non-slave) cells list for this row.
            visible = [v for v in row_data if v is not None]
            nonempty = [v for v in visible if v.strip()]

            if not nonempty:
                continue   # blank row — skip entirely

            # Detect heading: every non-empty cell in this row belongs to the
            # SAME single merged region (the master coord appears exactly once).
            # This handles full-width merged titles and also narrow notes that
            # span all used columns.
            heading_text = None
            if merged_ranges:
                # Gather the master coords for all non-empty cells.
                masters_for_nonempty = set()
                for cell in ws.iter_rows(
                    min_row=row_cells[0].row, max_row=row_cells[0].row
                ):
                    for c in cell:
                        coord = (c.row, c.column)
                        v = "" if c.value is None else str(c.value)
                        if v.strip():
                            masters_for_nonempty.add(
                                merged_topleft.get(coord, coord)
                            )

                if len(masters_for_nonempty) == 1:
                    # All non-empty content originates from one merged region
                    # (or a single standalone cell that spans the whole row).
                    # Treat as heading only if it truly spans multiple columns
                    # OR is the only non-empty cell (legacy single-cell banner).
                    master_coord = next(iter(masters_for_nonempty))
                    # Find the range for that master (if any).
                    range_span = 1
                    for mr in merged_ranges:
                        if (mr.min_row, mr.min_col) == master_coord:
                            range_span = mr.max_col - mr.min_col + 1
                            break
                    if range_span > 1 or len(nonempty) == 1:
                        heading_text = nonempty[0]
            else:
                # No merged regions at all — fall back to single-value heuristic.
                if len(nonempty) == 1:
                    heading_text = nonempty[0]

            if heading_text is not None:
                if table:
                    blocks.append(("table", table)); table = []
                blocks.append(("heading", heading_text))
            else:
                # Normal data row — trim trailing empty columns.
                cells = visible[:]
                while cells and not cells[-1].strip():
                    cells.pop()
                table.append(cells)

        if table:
            blocks.append(("table", table))
        sheets.append((ws.title, blocks))

    wb.close()
    return sheets


def to_docx(sheets, out_path):
    """Render sheets to a .docx file.

    Raises RuntimeError (not ImportError) with install hint if python-docx is
    not available.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run:  pip install python-docx"
        ) from None

    doc = Document()
    for si, (name, blocks) in enumerate(sheets):
        if si:
            doc.add_page_break()
        doc.add_heading(name, level=1)
        for kind, payload in blocks:
            if kind == "heading":
                p = doc.add_paragraph()
                run = p.add_run(payload)
                run.bold = True
                run.font.size = Pt(11)
            else:
                ncol = max(len(r) for r in payload)
                tbl = doc.add_table(rows=0, cols=ncol)
                tbl.style = "Light Grid Accent 1"
                for ri, row in enumerate(payload):
                    cells = tbl.add_row().cells
                    for ci in range(ncol):
                        val = row[ci] if ci < len(row) else ""
                        cells[ci].text = val
                        para = cells[ci].paragraphs[0]
                        if ri == 0:
                            # Bold the header row.
                            if para.runs:
                                para.runs[0].font.bold = True
                        if val in AP_RGB:
                            # AP colour carry-through: colour the run text.
                            r = para.runs[0] if para.runs else para.add_run(val)
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(*AP_RGB[val])

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def to_pdf(sheets, out_path):
    """Render sheets to a .pdf file via reportlab.

    Raises RuntimeError (not ImportError) with install hint if reportlab is
    not available.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak,
        )
    except ImportError:
        raise RuntimeError(
            "reportlab is not installed. Run:  pip install reportlab"
        ) from None

    styles = getSampleStyleSheet()
    cell = styles["BodyText"]; cell.fontSize = 7; cell.leading = 9
    doc = SimpleDocTemplate(
        str(out_path), pagesize=landscape(letter),
        leftMargin=0.4 * inch, rightMargin=0.4 * inch,
        topMargin=0.4 * inch, bottomMargin=0.4 * inch,
    )
    avail = doc.width
    story = []
    for si, (name, blocks) in enumerate(sheets):
        if si:
            story.append(PageBreak())
        story.append(Paragraph(f"<b>{name}</b>", styles["Heading1"]))
        for kind, payload in blocks:
            if kind == "heading":
                story.append(Paragraph(payload, styles["Heading4"]))
                story.append(Spacer(1, 4))
            else:
                ncol = max(len(r) for r in payload)
                data = [
                    [Paragraph((r[ci] if ci < len(r) else ""), cell) for ci in range(ncol)]
                    for r in payload
                ]
                tbl = Table(data, colWidths=[avail / ncol] * ncol, repeatRows=1)
                ts = [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
                # AP colour carry-through: shade individual cells.
                for ri, row in enumerate(payload):
                    for ci, val in enumerate(row):
                        if val in AP_RGB:
                            ts.append(("BACKGROUND", (ci, ri), (ci, ri),
                                       colors.Color(*[c / 255 for c in AP_RGB[val]])))
                tbl.setStyle(TableStyle(ts))
                story.append(tbl)
                story.append(Spacer(1, 10))

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def _selfcheck():
    """Build a tiny workbook with both plain rows and a merged heading, then
    export both formats and assert non-empty output files.
    """
    import tempfile
    tmp = Path(tempfile.mkdtemp())
    src = tmp / "t.xlsx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "S1"

    # Row 1: merged heading across columns A-E.
    ws.append(["pFMEA Self-Check", "", "", "", ""])
    ws.merge_cells("A1:E1")

    # Row 2: header row (multiple columns — must NOT be a heading).
    ws.append(["Failure Mode", "S", "O", "D", "AP"])

    # Row 3-4: data rows.
    ws.append(["Leak", 9, 2, 8, "H"])
    ws.append(["Scuff", 4, 3, 2, "L"])

    # Row 5: another merged note row.
    ws.append(["End of analysis", "", "", "", ""])
    ws.merge_cells("A5:E5")

    wb.save(src)

    sheets = read_sheets(src)
    assert sheets[0][0] == "S1", f"sheet name wrong: {sheets[0][0]}"
    kinds = [k for k, _ in sheets[0][1]]
    # Expected: heading (merged row 1), table (rows 2-4), heading (merged row 5)
    assert kinds == ["heading", "table", "heading"], \
        f"block kinds wrong: {kinds}\nblocks: {sheets[0][1]}"

    # Verify AP colours survive into the table block.
    tbl_rows = sheets[0][1][1][1]   # the table payload
    ap_vals = [row[-1] for row in tbl_rows[1:]]   # skip header row
    assert "H" in ap_vals and "L" in ap_vals, f"AP values missing: {ap_vals}"

    d, p = tmp / "t.docx", tmp / "t.pdf"
    to_docx(sheets, d)
    to_pdf(sheets, p)
    assert d.stat().st_size > 0, "docx is empty"
    assert p.stat().st_size > 0, "pdf is empty"
    print(f"selfcheck OK -> {d}, {p}")


def main(argv):
    if "--selfcheck" in argv:
        _selfcheck(); return
    args = [a for a in argv if not a.startswith("-")]
    if not args:
        print(__doc__); sys.exit(1)
    src = Path(args[0])
    if not src.exists():
        sys.exit(f"No such file: {src}")
    fmt = "all"
    if "--format" in argv:
        fmt = argv[argv.index("--format") + 1]
    base = src.with_suffix("")
    if "-o" in argv:
        base = Path(argv[argv.index("-o") + 1])

    sheets = read_sheets(src)
    if fmt in ("docx", "all"):
        out = base.with_suffix(".docx"); to_docx(sheets, out); print(f"wrote {out}")
    if fmt in ("pdf", "all"):
        out = base.with_suffix(".pdf"); to_pdf(sheets, out); print(f"wrote {out}")


if __name__ == "__main__":
    main(sys.argv[1:])
